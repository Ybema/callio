#!/usr/bin/env python3
"""
Word Export Module for Review Reports
Generates professional Word documents with proper formatting and branding.
"""

import json
import datetime
from pathlib import Path
from typing import Dict, Any, List, Optional
from docx import Document
from docx.shared import Inches, Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.style import WD_STYLE_TYPE
from docx.oxml.shared import OxmlElement, qn


class ReviewWordExporter:
    """Export review results to professional Word documents."""
    
    def __init__(self):
        self.doc = None
        self.verdana_font = "Verdana"
    
    def _find_historical_results(self, current_session_id: str, results_dir: Path) -> List[Dict[str, Any]]:
        """Find and load the 3 most recent historical review results."""
        if not results_dir.exists():
            return []
        
        # Find all JSON result files, excluding the current one
        json_files = []
        for file_path in results_dir.glob("lfa_review_result_*.json"):
            if current_session_id not in file_path.name:
                json_files.append(file_path)
        
        # Sort by modification time (newest first) and take the 3 most recent
        json_files.sort(key=lambda x: x.stat().st_mtime, reverse=True)
        historical_results = []
        
        for file_path in json_files[:3]:  # Take 3 most recent
            try:
                with open(file_path, 'r', encoding='utf-8') as f:
                    data = json.load(f)
                    # Extract version info from filename
                    version_timestamp = file_path.stem.split('_')[-1]  # Extract timestamp
                    data['version_label'] = f"v{version_timestamp[-4:]}"  # Use last 4 digits (MMSS)
                    historical_results.append(data)
            except Exception as e:
                continue  # Skip corrupted files
        
        # Reverse to show oldest to newest (left to right in table)
        return list(reversed(historical_results))
    
    def _get_score_color(self, score: float) -> tuple:
        """Get color for score based on 0-5 scale: red→orange→yellow→light green→green."""
        if score == 0:
            return (128, 128, 128)  # Gray for N/A or 0
        elif score <= 1.0:
            return (220, 53, 69)   # Red
        elif score <= 2.0:
            return (255, 133, 27)  # Orange  
        elif score <= 3.0:
            return (255, 193, 7)   # Yellow
        elif score <= 4.0:
            return (40, 167, 69)   # Light Green
        else:  # score > 4.0
            return (25, 135, 84)   # Green
        
    def export_review_to_word(self, review_data: Dict[str, Any], output_path: str) -> str:
        """
        Export review results to a Word document.
        
        Args:
            review_data: The review results dictionary
            output_path: Path where to save the Word document
            
        Returns:
            Path to the created Word document
        """
        self.doc = Document()
        self._setup_document_styles()
        
        # Add metadata header
        self._add_metadata_section(review_data)
        
        # Add score explanation (right after header)
        self._add_score_explanation(review_data)
        
        # Add executive summary
        self._add_executive_summary(review_data)
        
        # Add progress tracking table
        self._add_progress_table(review_data, Path(output_path).parent)
        
        # Add page break before detailed findings
        self.doc.add_page_break()
        
        # Add detailed findings
        self._add_detailed_findings(review_data)
        
        # Add footer with branding
        self._add_footer()
        
        # Save document
        self.doc.save(output_path)
        return output_path
    
    def _setup_document_styles(self):
        """Set up document styles with Verdana font."""
        # Set default font for the document
        style = self.doc.styles['Normal']
        font = style.font
        font.name = self.verdana_font
        font.size = Pt(10)
        
        # Create custom styles
        self._create_heading_style("Heading 1", 16, True)
        self._create_heading_style("Heading 2", 14, True)
        self._create_heading_style("Heading 3", 12, True)
        self._create_heading_style("Metadata", 9, False)
        self._create_heading_style("Score", 11, True)
        self._create_heading_style("Evidence", 10, True)
        self._create_heading_style("Gaps", 10, True)
        self._create_heading_style("Fixes", 10, True)
    
    def _create_heading_style(self, name: str, size: int, bold: bool):
        """Create a custom heading style."""
        try:
            style = self.doc.styles.add_style(name, WD_STYLE_TYPE.PARAGRAPH)
        except ValueError:
            # Style already exists
            style = self.doc.styles[name]
        
        font = style.font
        font.name = self.verdana_font
        font.size = Pt(size)
        font.bold = bold
        
        if name.startswith("Heading"):
            style.paragraph_format.space_before = Pt(12)
            style.paragraph_format.space_after = Pt(6)
    
    def _add_metadata_section(self, review_data: Dict[str, Any]):
        """Add metadata header section."""
        # Title with project name
        meta = review_data.get('meta', {})
        project_name = meta.get('project_name', 'Project')
        title = self.doc.add_heading(f'{project_name} - Horizon Europe Proposal Review Report', 0)
        title.alignment = WD_ALIGN_PARAGRAPH.CENTER
        
        # Metadata table
        meta = review_data.get('meta', {})
        scores = review_data.get('scores', {})
        
        # Populate metadata
        metadata_items = [
            ("Project", project_name),
            ("Review Phase", "Phase A - Logic Framework Analysis"),
            ("Report Generated", datetime.datetime.now().strftime("%Y-%m-%d %H:%M:%S")),
            ("Engine Version", meta.get('engine_version', 'N/A')),
            ("AI Model", meta.get('model', 'N/A')),
            ("Session ID", meta.get('timestamp', 'N/A')),
            ("Total Score", f"{scores.get('total', 0):.2f}/5.0 ({scores.get('band', 'N/A')})"),
        ]
        score_labels = {
            "CA": "Call Alignment",
            "LC": "Internal Consistency",
            "CQ": "Content Quality",
        }
        default_weights = {"CA": 0.30, "LC": 0.50, "CQ": 0.20}
        for key in ["CA", "LC", "CQ"]:
            category_data = scores.get(key)
            if isinstance(category_data, dict):
                subtotal = float(category_data.get("subtotal", 0.0))
                weight = float(category_data.get("weight", default_weights[key]))
                metadata_items.append(
                    (score_labels[key], f"{subtotal:.2f}/5.0 ({int(weight * 100)}% weight)")
                )

        table = self.doc.add_table(rows=len(metadata_items), cols=2)
        table.style = 'Table Grid'
        
        for i, (label, value) in enumerate(metadata_items):
            table.cell(i, 0).text = label
            table.cell(i, 1).text = str(value)
            
            # Style the cells with smaller font for compactness
            for cell in [table.cell(i, 0), table.cell(i, 1)]:
                for paragraph in cell.paragraphs:
                    paragraph.style = 'Metadata'
                    for run in paragraph.runs:
                        run.font.size = Pt(9)  # Smaller font for compactness
        
        # Usage information
        cost_tracking = meta.get('cost_tracking', {})
        if cost_tracking:
            usage_para = self.doc.add_paragraph()
            usage_para.add_run("Usage Analysis: ").bold = True
            usage_para.add_run(f"Tokens: {cost_tracking.get('total_tokens', 0):,} | ")
            usage_para.add_run(f"API Calls: {cost_tracking.get('api_calls', 0)} | ")
            usage_para.add_run(f"Provider: {', '.join(cost_tracking.get('provider_usage', {}).keys())}")
    
    def _add_score_explanation(self, review_data: Dict[str, Any]):
        """Add compact score explanation section."""
        self.doc.add_heading('Score Explanation', 1)
        
        # Compact explanation in a single paragraph
        explanation = self.doc.add_paragraph()
        exp_run = explanation.add_run("Scoring: ")
        exp_run.bold = True
        exp_run.font.size = Pt(9)
        
        score_text = "5=Excellent, 4=Good, 3=Adequate, 2=Needs Work, 1=Poor. "
        score_run = explanation.add_run(score_text)
        score_run.font.size = Pt(9)
        
        scores = review_data.get("scores", {})
        default_weights = {"CA": 0.30, "LC": 0.50, "CQ": 0.20}
        labels = {"CA": "CA", "LC": "LC", "CQ": "CQ"}
        weighted_parts = []
        for key in ["CA", "LC", "CQ"]:
            category_data = scores.get(key)
            if isinstance(category_data, dict):
                weight = float(category_data.get("weight", default_weights[key]))
                weighted_parts.append(f"{labels[key]}({int(weight * 100)}%)")

        if weighted_parts:
            formula = "Total = " + " + ".join(weighted_parts)
        else:
            formula = "Total = Weighted average of available criteria"

        calc_run = explanation.add_run(formula)
        calc_run.bold = True
        calc_run.font.size = Pt(9)

    def _add_executive_summary(self, review_data: Dict[str, Any]):
        """Add executive summary section."""
        self.doc.add_heading('Executive Summary', 1)
        
        scores = review_data.get('scores', {})
        total_score = scores.get('total', 0)
        band = scores.get('band', 'N/A')
        
        # Get individual scores for categories that were actually evaluated
        ca_data = scores.get("CA")
        lc_data = scores.get("LC")
        cq_data = scores.get("CQ")

        ca_score = float(ca_data.get("subtotal", 0.0)) if isinstance(ca_data, dict) else None
        lc_score = float(lc_data.get("subtotal", 0.0)) if isinstance(lc_data, dict) else None
        cq_score = float(cq_data.get("subtotal", 0.0)) if isinstance(cq_data, dict) else None
        
        summary_para = self.doc.add_paragraph()
        
        # Check for missing scores
        missing_scores = []
        if ca_score is not None and ca_score == 0.0:
            missing_scores.append("Call Alignment")
        if lc_score is not None and lc_score == 0.0:
            missing_scores.append("Logic Consistency")
        if cq_score is not None and cq_score == 0.0:
            missing_scores.append("Content Quality")
        
        if missing_scores:
            summary_para.add_run(f"Note: {', '.join(missing_scores)} evaluation failed due to technical issues. ").italic = True
        
        summary_para.add_run(f"This proposal received an overall score of {total_score:.2f}/5.0, placing it in the '{band}' category. ")
        
        if ca_score is not None:
            if ca_score >= 4.0:
                summary_para.add_run("The proposal shows strong alignment with the call requirements. ")
            elif ca_score >= 3.0:
                summary_para.add_run("The proposal demonstrates adequate alignment with the call requirements, with room for improvement. ")
            else:
                summary_para.add_run("The proposal requires significant improvements in call alignment. ")
        
        if lc_score is not None:
            if lc_score >= 4.0:
                summary_para.add_run("The logical framework is well-structured and internally consistent. ")
            elif lc_score >= 3.0:
                summary_para.add_run("The logical framework shows good structure but has some consistency issues. ")
            else:
                summary_para.add_run("The logical framework needs substantial restructuring for better consistency. ")
        
        if cq_score is not None:
            if cq_score >= 4.0:
                summary_para.add_run("The content quality is high with clear, well-supported arguments. ")
            elif cq_score >= 3.0:
                summary_para.add_run("The content quality is adequate but could benefit from more specific evidence and clearer language. ")
            else:
                summary_para.add_run("The content quality needs significant improvement in clarity and evidence. ")
        
        summary_para.add_run("Detailed findings and recommendations are provided in the following sections.")
    
    def _add_progress_table(self, review_data: Dict[str, Any], results_dir: Path):
        """Add progress tracking table showing score evolution across versions."""
        # Get current session ID for exclusion
        current_session_id = review_data.get('meta', {}).get('timestamp', '').replace('-', '').replace(':', '')
        
        # Load historical results
        historical_results = self._find_historical_results(current_session_id, results_dir)
        
        # Only show table if we have historical data
        if not historical_results:
            return
        
        self.doc.add_heading('Progress Tracking', 1)
        
        # Create table with columns: Criterion + historical versions + current
        num_versions = len(historical_results) + 1  # +1 for current
        table = self.doc.add_table(rows=1, cols=num_versions + 1)  # +1 for criterion column
        table.style = 'Table Grid'
        
        # Set header row
        header_cells = table.rows[0].cells
        header_cells[0].text = "Criterion"
        
        # Add version headers
        for i, hist_result in enumerate(historical_results):
            header_cells[i + 1].text = hist_result.get('version_label', f'v{i+1}')
        header_cells[-1].text = "Current"
        
        # Style header row with smaller font
        for cell in header_cells:
            for paragraph in cell.paragraphs:
                for run in paragraph.runs:
                    run.font.size = Pt(9)
                    run.font.name = self.verdana_font
                    run.bold = True
        
        # Get all criteria from current results
        current_scores = review_data.get('scores', {})
        criteria_list = []
        
        # Add individual criteria scores
        for category in ['CA', 'LC', 'CQ']:
            if category in current_scores:
                category_scores = current_scores[category]
                for key, value in category_scores.items():
                    if key not in ['weight', 'subtotal'] and isinstance(value, (int, float)):
                        criteria_list.append(key)
        
        # Add rows for each criterion
        for criterion in sorted(criteria_list):
            row = table.add_row()
            cells = row.cells
            
            # Criterion name
            cells[0].text = criterion
            cells[0].paragraphs[0].runs[0].font.size = Pt(9)
            cells[0].paragraphs[0].runs[0].font.name = self.verdana_font
            
            # Historical scores
            for i, hist_result in enumerate(historical_results):
                hist_scores = hist_result.get('scores', {})
                score = self._get_criterion_score(hist_scores, criterion)
                
                # Set cell text and color
                if score is None:
                    cells[i + 1].text = "N/A"
                    bg_color = (128, 128, 128)  # Gray
                else:
                    cells[i + 1].text = f"{score:.1f}"
                    bg_color = self._get_score_color(score)
                
                # Apply formatting
                self._format_score_cell(cells[i + 1], bg_color)
            
            # Current score
            current_score = self._get_criterion_score(current_scores, criterion)
            if current_score is None:
                cells[-1].text = "N/A"
                bg_color = (128, 128, 128)
            else:
                cells[-1].text = f"{current_score:.1f}"
                bg_color = self._get_score_color(current_score)
            
            self._format_score_cell(cells[-1], bg_color)
    
    def _get_criterion_score(self, scores_dict: Dict[str, Any], criterion: str) -> Optional[float]:
        """Extract score for a specific criterion from scores dictionary."""
        # Determine category (CA, LC, CQ) from criterion name
        category = criterion[:2]
        if category in scores_dict and isinstance(scores_dict[category], dict):
            return scores_dict[category].get(criterion, None)
        return None
    
    def _format_score_cell(self, cell, bg_color: tuple):
        """Apply formatting to a score cell with background color."""
        # Set text formatting
        for paragraph in cell.paragraphs:
            for run in paragraph.runs:
                run.font.size = Pt(9)
                run.font.name = self.verdana_font
                run.bold = True
                # Use white text on dark backgrounds, black on light
                if bg_color[0] + bg_color[1] + bg_color[2] < 400:
                    run.font.color.rgb = RGBColor(255, 255, 255)  # White
                else:
                    run.font.color.rgb = RGBColor(0, 0, 0)  # Black
        
        # Set background color
        cell_xml = cell._tc
        cell_pr = cell_xml.get_or_add_tcPr()
        shading = OxmlElement('w:shd')
        # Convert RGBColor to hex string
        hex_color = f"{bg_color[0]:02x}{bg_color[1]:02x}{bg_color[2]:02x}"
        shading.set(qn('w:fill'), hex_color)
        cell_pr.append(shading)
    
    def _add_detailed_findings(self, review_data: Dict[str, Any]):
        """Add detailed findings section."""
        self.doc.add_heading('Detailed Findings', 1)
        
        findings = review_data.get('findings', {})
        
        # Group findings by category
        ca_findings = {k: v for k, v in findings.items() if k.startswith('CA') and not k.endswith('_ERROR')}
        lc_findings = {k: v for k, v in findings.items() if k.startswith('LC') and not k.endswith('_ERROR')}
        cq_findings = {k: v for k, v in findings.items() if k.startswith('CQ') and not k.endswith('_ERROR')}
        
        # Add Call Alignment section
        if ca_findings:
            self.doc.add_heading('Call Alignment (CA)', 2)
            self._add_criteria_section(ca_findings, review_data.get('scores', {}).get('CA', {}), "CA")
        
        # Add Logic Consistency section
        if lc_findings:
            self.doc.add_heading('Logic Consistency (LC)', 2)
            self._add_criteria_section(lc_findings, review_data.get('scores', {}).get('LC', {}), "LC")
        
        # Add Content Quality section
        if cq_findings:
            self.doc.add_heading('Content Quality (CQ)', 2)
            self._add_criteria_section(cq_findings, review_data.get('scores', {}).get('CQ', {}), "CQ")
    
    def _add_criteria_section(self, findings: Dict[str, Any], scores: Dict[str, Any], criteria_type: str):
        """Add a criteria section with detailed findings."""
        # Define criteria names and explanations
        criteria_names = {
            "CA1": "Objectives alignment",
            "CA2": "Scope fit", 
            "CA3": "Outcomes/Impacts alignment",
            "CA4": "Evaluation coverage",
            "CA5": "LFA-Call alignment completeness",
            "CA6": "Terminology/definitions alignment",
            "LC1": "Logical Flow",
            "LC2": "Measurable Outcomes", 
            "LC3": "Activity-Outcome Linkage",
            "LC4": "Implementation Feasibility",
            "CQ1": "Clarity & Specificity",
            "CQ2": "Actionable Content",
            "CQ3": "Professional Presentation"
        }
        
        criteria_explanations = {
            "CA": {
                "CA1": "How well the LFA objectives align with the call's main aims and requirements",
                "CA2": "Whether the LFA activities, target populations, and geographies fit within the call's scope",
                "CA3": "How well the LFA outcomes address the call's expected impacts and success criteria",
                "CA4": "Whether the LFA addresses the call's evaluation criteria (Excellence, Impact, Implementation)",
                "CA5": "Overall completeness of alignment between LFA elements and call requirements",
                "CA6": "Consistency in using call-specific terminology and definitions"
            },
            "LC": {
                "LC1": "Clear progression from Goal → Purpose → Outcomes → Activities",
                "LC2": "Specific, quantifiable outcomes with clear indicators and targets",
                "LC3": "Clear connection between activities and outcomes they enable",
                "LC4": "Activities are realistic, achievable, and well-defined"
            },
            "CQ": {
                "CQ1": "Clear, specific language with concrete examples and quantified claims",
                "CQ2": "Content that can be implemented, measured, and tracked with clear next steps",
                "CQ3": "Well-structured, professional writing with clear organization and minimal redundancy"
            }
        }
        
        criteria_type_weights = {
            "CA": 0.30,
            "LC": 0.40,
            "CQ": 0.20
        }
        
        for criterion, finding in findings.items():
            if criterion in scores:
                score = scores[criterion]
                explanation = criteria_explanations.get(criteria_type, {}).get(criterion, f"{criterion}: Evaluation criterion")
                
                # Use full criterion name with explanation
                criterion_name = criteria_names.get(criterion, criterion)
                self.doc.add_heading(f'{criterion}: {criterion_name}', 3)
                
                # Score and weight
                weight_percent = int(criteria_type_weights.get(criteria_type, 0.2) * 100)
                score_para = self.doc.add_paragraph()
                score_para.add_run(f'Score: {score}/5.0').bold = True
                score_para.add_run(f' | Weight: {weight_percent}%')
                
                # What we're evaluating
                eval_para = self.doc.add_paragraph()
                eval_para.add_run("What we're evaluating: ").bold = True
                eval_para.add_run(explanation)
                
                # Add strengths (evidence) - now narrative format
                if 'evidence' in finding and finding['evidence']:
                    strengths_para = self.doc.add_paragraph()
                    strengths_para.add_run('Strengths:').bold = True
                    for evidence in finding['evidence']:
                        if isinstance(evidence, dict):
                            quote = evidence.get('quote', '')
                            loc = evidence.get('loc', '')
                            if quote:
                                para = self.doc.add_paragraph(f'• "{quote}"', style='Normal')
                                if loc:
                                    para.add_run(f' — {loc}').italic = True
                        else:
                            # Handle narrative format - no bullets, just paragraph text
                            para = self.doc.add_paragraph(evidence, style='Normal')
                
                # Add areas for improvement (gaps) - now narrative format
                if 'gaps' in finding and finding['gaps']:
                    gaps_para = self.doc.add_paragraph()
                    gaps_para.add_run('Areas for Improvement:').bold = True
                    for gap in finding['gaps']:
                        # Handle narrative format - no bullets, just paragraph text
                        para = self.doc.add_paragraph(gap, style='Normal')
                
                # Add suggestions (fixes) - now narrative format
                if 'fixes' in finding and finding['fixes']:
                    suggestions_para = self.doc.add_paragraph()
                    suggestions_para.add_run('Suggestions to straight copy-paste into the document:').bold = True
                    for fix in finding['fixes']:
                        # Handle narrative format - no bullets, just paragraph text
                        para = self.doc.add_paragraph(fix, style='Normal')
                
                # Add details for special criteria
                if 'details' in finding:
                    self.doc.add_paragraph('Technical Details:', style='Evidence')
                    details = finding['details']
                    if isinstance(details, dict):
                        for key, value in details.items():
                            if isinstance(value, list) and value:
                                self.doc.add_paragraph(f'{key.replace("_", " ").title()}:', style='Evidence')
                                for item in value:
                                    self.doc.add_paragraph(f'• {item}', style='Normal')
                            elif value:
                                self.doc.add_paragraph(f'{key.replace("_", " ").title()}: {value}', style='Normal')
    
    def _add_footer(self):
        """Add footer with Sustainovate branding."""
        # Add a page break before footer
        self.doc.add_page_break()
        
        # Add footer section
        footer_heading = self.doc.add_heading('Report Information', 1)
        footer_heading.alignment = WD_ALIGN_PARAGRAPH.CENTER
        
        # Add branding paragraph
        branding_para = self.doc.add_paragraph()
        branding_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
        branding_para.add_run('This report has been generated using ').font.size = Pt(10)
        branding_para.add_run('Sustainovate AI Technology').font.size = Pt(10)
        branding_para.add_run('.').font.size = Pt(10)
        
        # Add disclaimer
        disclaimer_para = self.doc.add_paragraph()
        disclaimer_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
        disclaimer_para.add_run('This automated analysis is provided for guidance purposes. ').font.size = Pt(9)
        disclaimer_para.add_run('Please review all findings and recommendations carefully before making final decisions.').font.size = Pt(9)
        
        # Add generation timestamp
        timestamp_para = self.doc.add_paragraph()
        timestamp_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
        timestamp_para.add_run(f'Report generated on {datetime.datetime.now().strftime("%Y-%m-%d at %H:%M:%S")}').font.size = Pt(8)


def export_review_to_word(review_data: Dict[str, Any], output_path: str) -> str:
    """
    Convenience function to export review results to Word.
    
    Args:
        review_data: The review results dictionary
        output_path: Path where to save the Word document
        
    Returns:
        Path to the created Word document
    """
    exporter = ReviewWordExporter()
    return exporter.export_review_to_word(review_data, output_path)


if __name__ == "__main__":
    # Test the exporter
    test_data = {
        "meta": {
            "engine_version": "0.2.0",
            "model": "gpt-4o-mini",
            "timestamp": "2025-09-24T22:50:58"
        },
        "scores": {
            "total": 2.64,
            "band": "Needs Work",
            "CA": {"subtotal": 3.0},
            "LC": {"subtotal": 2.23},
            "CQ": {"subtotal": 3.12}
        },
        "findings": {
            "CA1": {
                "evidence": [{"quote": "Test evidence", "loc": "Test location"}],
                "gaps": ["Test gap"],
                "fixes": ["Test fix"]
            }
        }
    }
    
    output_path = "test_review_report.docx"
    result_path = export_review_to_word(test_data, output_path)
    print(f"Test Word document created: {result_path}")
