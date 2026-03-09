#!/usr/bin/env python3
"""
Word Document Export for Proposal Framework
Converts complete markdown analysis files to professional Word documents
"""

import os
import sys
import re
import yaml
from docx import Document
from docx.shared import Inches, RGBColor, Pt
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.shared import OxmlElement, qn
from docx.enum.style import WD_STYLE_TYPE

def load_formatting_config():
    """Load formatting configuration from YAML file"""
    # Try multiple possible paths for the config file
    possible_paths = [
        "templates/output_templates/word_formatting_definitions.yaml",
        "project framework/templates/output_templates/word_formatting_definitions.yaml",
        "output/templates/word_formatting_definitions.yaml"  # Fallback to original location
    ]
    
    for config_path in possible_paths:
        try:
            with open(config_path, 'r', encoding='utf-8') as f:
                print(f"📄 Using formatting config: {config_path}")
                return yaml.safe_load(f)
        except FileNotFoundError:
            continue
    
    print(f"⚠️  Warning: No formatting config found, using default formatting")
    return get_default_config()

def get_default_config():
    """Default configuration if YAML file not found"""
    return {
        'score_colors': {
            'exceptional': {'background': 'D4E6F1', 'text': '1B4F72'},
            'excellent': {'background': 'D5F4E6', 'text': '196F3D'},
            'good': {'background': 'FCF3CF', 'text': 'B7950B'},
            'fair': {'background': 'FADBD8', 'text': 'A93226'},
            'poor': {'background': 'FADBD8', 'text': 'C0392B'}
        }
    }

def hex_to_rgb(hex_color):
    """Convert hex color to RGB tuple"""
    hex_color = hex_color.lstrip('#')
    return tuple(int(hex_color[i:i+2], 16) for i in (0, 2, 4))

def get_score_category(score):
    """Determine score category for formatting"""
    if score >= 90:
        return 'exceptional'
    elif score >= 85:
        return 'excellent'
    elif score >= 75:
        return 'good'
    elif score >= 65:
        return 'fair'
    else:
        return 'poor'

def apply_cell_formatting(cell, score=None, status=None, config=None):
    """Apply formatting to table cell based on score or status"""
    if not config:
        return
        
    category = None
    if score is not None:
        category = get_score_category(score)
    elif status:
        status_map = {
            'EXCEPTIONAL': 'exceptional',
            'EXCELLENT': 'excellent', 
            'GOOD': 'good',
            'FAIR': 'fair',
            'POOR': 'poor'
        }
        category = status_map.get(status.upper())
    
    if category and category in config.get('score_colors', {}):
        color_config = config['score_colors'][category]
        # Set background color
        if 'background' in color_config:
            # Clear any existing shading first
            tc_pr = cell._tc.get_or_add_tcPr()
            # Remove existing shading elements
            for shading in tc_pr.xpath('.//w:shd'):
                tc_pr.remove(shading)
            
            # Apply new shading
            shading_elm = OxmlElement('w:shd')
            shading_elm.set(qn('w:fill'), color_config['background'])
            shading_elm.set(qn('w:val'), 'clear')  # Ensure it's visible
            tc_pr.append(shading_elm)
        
        # Set text color and bold
        if 'text' in color_config:
            for paragraph in cell.paragraphs:
                for run in paragraph.runs:
                    text_color = hex_to_rgb(color_config['text'])
                    run.font.color.rgb = RGBColor(*text_color)
                    run.font.bold = True

def setup_document_styles(doc, config):
    """Setup custom styles for the document using YAML configuration"""
    styles = doc.styles
    doc_settings = config.get('document_settings', {})
    text_styles = config.get('text_styles', {})
    
    font_family = doc_settings.get('font_family', 'Aptos')
    base_font_size = doc_settings.get('font_size', 10)
    
    # Title style from YAML configuration
    if 'Title Custom' not in [s.name for s in styles]:
        title_style = styles.add_style('Title Custom', WD_STYLE_TYPE.PARAGRAPH)
        title_config = text_styles.get('title', {})
        
        title_style.font.name = font_family
        title_style.font.size = Pt(title_config.get('font_size', 16))
        title_style.font.bold = title_config.get('bold', True)
        
        # Convert hex color to RGB
        title_color = title_config.get('color', '2E86C1')
        title_rgb = hex_to_rgb(title_color)
        title_style.font.color.rgb = RGBColor(*title_rgb)
        
        title_style.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.CENTER
        title_style.paragraph_format.space_before = Pt(0)
        title_style.paragraph_format.space_after = Pt(6)
        title_style.paragraph_format.line_spacing = 1.0
    
    # Heading 1 style from YAML configuration
    if 'Heading 1 Custom' not in [s.name for s in styles]:
        h1_style = styles.add_style('Heading 1 Custom', WD_STYLE_TYPE.PARAGRAPH)
        h1_config = text_styles.get('section_header', {})
        
        h1_style.font.name = font_family
        h1_style.font.size = Pt(h1_config.get('font_size', 14))
        h1_style.font.bold = h1_config.get('bold', True)
        
        # Convert hex color to RGB
        h1_color = h1_config.get('color', '1B4F72')
        h1_rgb = hex_to_rgb(h1_color)
        h1_style.font.color.rgb = RGBColor(*h1_rgb)
        
        h1_style.paragraph_format.space_before = Pt(8)
        h1_style.paragraph_format.space_after = Pt(3)
        h1_style.paragraph_format.line_spacing = 1.0
    
    # Heading 2 style from YAML configuration  
    if 'Heading 2 Custom' not in [s.name for s in styles]:
        h2_style = styles.add_style('Heading 2 Custom', WD_STYLE_TYPE.PARAGRAPH)
        h2_config = text_styles.get('subsection_header', {})
        
        h2_style.font.name = font_family
        h2_style.font.size = Pt(h2_config.get('font_size', 12))
        h2_style.font.bold = h2_config.get('bold', True)
        
        # Convert hex color to RGB
        h2_color = h2_config.get('color', '5D6D7E')
        h2_rgb = hex_to_rgb(h2_color)
        h2_style.font.color.rgb = RGBColor(*h2_rgb)
        
        h2_style.paragraph_format.space_before = Pt(6)
        h2_style.paragraph_format.space_after = Pt(2)
        h2_style.paragraph_format.line_spacing = 1.0
    
    # Heading 3 style from YAML configuration
    if 'Heading 3 Custom' not in [s.name for s in styles]:
        h3_style = styles.add_style('Heading 3 Custom', WD_STYLE_TYPE.PARAGRAPH)
        h3_config = text_styles.get('detail_header', {})
        
        h3_style.font.name = font_family
        h3_style.font.size = Pt(h3_config.get('font_size', base_font_size + 1))
        h3_style.font.bold = h3_config.get('bold', True)
        
        # Convert hex color to RGB
        h3_color = h3_config.get('color', '5D6D7E')
        h3_rgb = hex_to_rgb(h3_color)
        h3_style.font.color.rgb = RGBColor(*h3_rgb)
        
        h3_style.paragraph_format.space_before = Pt(4)
        h3_style.paragraph_format.space_after = Pt(1)
        h3_style.paragraph_format.line_spacing = 1.0
    
    # Normal paragraph style from YAML configuration
    if 'Normal Compact' not in [s.name for s in styles]:
        normal_style = styles.add_style('Normal Compact', WD_STYLE_TYPE.PARAGRAPH)
        normal_config = text_styles.get('body_text', {})
        
        normal_style.font.name = font_family
        normal_style.font.size = Pt(normal_config.get('font_size', base_font_size))
        normal_style.paragraph_format.space_before = Pt(0)
        normal_style.paragraph_format.space_after = Pt(2)
        normal_style.paragraph_format.line_spacing = 1.0

def parse_markdown_table(table_text):
    """Parse markdown table into structured data"""
    lines = [line.strip() for line in table_text.split('\n') if line.strip()]
    
    # Remove table separators (lines with |---|---|)
    content_lines = [line for line in lines if not re.match(r'^\|[\s\-\|]+\|$', line)]
    
    rows = []
    for line in content_lines:
        if line.startswith('|') and line.endswith('|'):
            # Split by | and clean up
            cells = [cell.strip() for cell in line[1:-1].split('|')]
            rows.append(cells)
    
    return rows

def assign_column_widths(table, table_data, table_type="default"):
    """Assign intelligent column widths based on content and table type"""
    if not table_data or len(table_data) < 1:
        return
    
    headers = table_data[0] if table_data else []
    total_cols = len(headers)
    
    # Define width ratios based on table type and content analysis
    if 'performance' in table_type.lower():
        # WP Performance Matrix - prioritize Priority Actions column
        widths = []
        for i, header in enumerate(headers):
            header_lower = header.lower()
            if 'priority' in header_lower or 'action' in header_lower:
                widths.append(3.5)  # Wide for detailed actions
            elif 'wp' in header_lower and len(header) < 5:
                widths.append(0.6)  # Compact for WP names
            elif 'status' in header_lower:
                widths.append(0.7)  # Compact for status
            elif '%' in header_lower or 'score' in header_lower:
                widths.append(0.8)  # Compact for percentages
            else:
                widths.append(1.0)  # Default width
    
    elif 'cross' in table_type.lower():
        # Cross-WP Matrix - prioritize Critical Gaps/Questions column
        widths = []
        for i, header in enumerate(headers):
            header_lower = header.lower()
            if 'gap' in header_lower or 'question' in header_lower or 'critical' in header_lower:
                widths.append(4.0)  # Very wide for detailed gaps/questions
            elif 'outcome' in header_lower:
                widths.append(1.5)  # Medium for outcome names
            elif 'wp' in header_lower or 'leading' in header_lower or 'supporting' in header_lower:
                widths.append(1.0)  # Compact for WP lists
            elif 'quality' in header_lower or '%' in header_lower:
                widths.append(0.8)  # Compact for percentages
            else:
                widths.append(1.2)  # Default width
    
    else:
        # Default table - analyze content length
        widths = []
        for i, header in enumerate(headers):
            # Analyze content length in this column
            avg_length = sum(len(str(row[i])) if i < len(row) else 0 for row in table_data) / len(table_data)
            if avg_length > 50:  # Long content
                widths.append(3.0)
            elif avg_length > 30:  # Medium content
                widths.append(2.0)
            elif avg_length > 15:  # Short content
                widths.append(1.2)
            else:  # Very short content
                widths.append(0.8)
    
    # Apply widths to table columns
    for i, width in enumerate(widths[:total_cols]):
        if i < len(table.columns):
            table.columns[i].width = Inches(width)

def create_formatted_table(doc, table_data, config, table_type="default"):
    """Create a formatted table from data with intelligent column widths"""
    if not table_data or len(table_data) < 2:
        return None
    
    # Create table
    table = doc.add_table(rows=len(table_data), cols=len(table_data[0]))
    table.style = 'Table Grid'
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    
    # Apply intelligent column widths
    assign_column_widths(table, table_data, table_type)
    
    # Set header background color based on table type
    header_color = '2E86C1'  # Default blue
    if 'cross_wp' in table_type.lower():
        header_color = '28B463'  # Green for cross-WP tables
    
    # Populate table
    for row_idx, row_data in enumerate(table_data):
        for col_idx, cell_data in enumerate(row_data):
            cell = table.cell(row_idx, col_idx)
            
            # Process bold and italic formatting in cell text first (before cleaning)
            if '*' in cell_data:
                # Clear cell and add formatted text
                cell.text = ""
                paragraph = cell.paragraphs[0]
                
                # First handle bold formatting (**text**)
                bold_parts = re.split(r'(\*\*[^*]*\*\*)', cell_data)
                
                # Then handle italic formatting (*text*)
                final_parts = []
                for bold_part in bold_parts:
                    if not (bold_part.startswith('**') and bold_part.endswith('**')):
                        # Split non-bold parts by italic markers
                        italic_parts = re.split(r'(\*[^*]*\*)', bold_part)
                        final_parts.extend(italic_parts)
                    else:
                        final_parts.append(bold_part)
                
                # Process all parts
                for part in final_parts:
                    if part.startswith('**') and part.endswith('**'):
                        # Bold text
                        bold_text = clean_text_content(part[2:-2])
                        if bold_text:
                            run = paragraph.add_run(bold_text)
                            run.bold = True
                    elif part.startswith('*') and part.endswith('*') and not part.startswith('**'):
                        # Italic text
                        italic_text = clean_text_content(part[1:-1])
                        if italic_text:
                            run = paragraph.add_run(italic_text)
                            run.italic = True
                    elif part:
                        # Regular text
                        clean_part = clean_text_content(part)
                        if clean_part:
                            run = paragraph.add_run(clean_part)
            else:
                cell.text = clean_text_content(cell_data)
            
            # Set compact paragraph formatting for all cells
            for paragraph in cell.paragraphs:
                paragraph.paragraph_format.space_before = Pt(1)
                paragraph.paragraph_format.space_after = Pt(1)
                paragraph.paragraph_format.line_spacing = 1.0
            
            # Apply formatting
            if row_idx == 0:  # Header row
                cell.paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER
                for paragraph in cell.paragraphs:
                    for run in paragraph.runs:
                        run.font.bold = True
                        run.font.color.rgb = RGBColor(255, 255, 255)
                        run.font.size = Pt(10)  # Smaller header font
                # Header background
                shading_elm = OxmlElement('w:shd')
                shading_elm.set(qn('w:fill'), header_color)
                cell._tc.get_or_add_tcPr().append(shading_elm)
            else:
                # Data rows - apply score-based formatting
                for paragraph in cell.paragraphs:
                    for run in paragraph.runs:
                        run.font.size = Pt(9)  # Smaller data font for compactness
                
                # Check if cell contains percentage (support decimal format like 85.0%)
                score_match = re.search(r'(\d+(?:\.\d+)?)%', cell_data)
                if score_match:
                    score = float(score_match.group(1))
                    apply_cell_formatting(cell, score=score, config=config)
                
                # Check if cell contains status
                status_words = ['EXCEPTIONAL', 'EXCELLENT', 'GOOD', 'FAIR', 'POOR']
                for status in status_words:
                    if status in cell_data.upper():
                        apply_cell_formatting(cell, status=status, config=config)
                        break
    
    return table

def clean_text_content(text):
    """Remove emojis, markdown artifacts and clean up text for Word document"""
    import re
    
    # Remove emojis (comprehensive Unicode ranges)
    emoji_pattern = re.compile("["
                               u"\U0001F600-\U0001F64F"  # emoticons
                               u"\U0001F300-\U0001F5FF"  # symbols & pictographs
                               u"\U0001F680-\U0001F6FF"  # transport & map
                               u"\U0001F1E0-\U0001F1FF"  # flags (iOS)
                               u"\U00002702-\U000027B0"  # dingbats
                               u"\U000024C2-\U0001F251"  # misc symbols
                               u"\U00002600-\U000026FF"  # miscellaneous symbols
                               u"\U0001F900-\U0001F9FF"  # supplemental symbols
                               "]+", flags=re.UNICODE)
    text = emoji_pattern.sub('', text)
    
    # Remove markdown formatting artifacts
    # Remove ** around text (but preserve the text for bold processing later)
    # This is handled in the calling functions, so we don't remove ** here
    
    # Remove other markdown artifacts
    text = re.sub(r'^\s*---\s*$', '', text)  # Remove horizontal rules
    text = re.sub(r'^\s*\*\s*$', '', text)   # Remove standalone asterisks
    text = re.sub(r'^\s*\+\s*$', '', text)   # Remove standalone plus signs
    text = re.sub(r'^\s*-\s*$', '', text)    # Remove standalone dashes
    
    # Clean up extra whitespace
    text = re.sub(r'\s+', ' ', text).strip()
    
    return text

def process_markdown_content(content, doc, config):
    """Process markdown content and convert to Word document"""
    lines = content.split('\n')
    i = 0
    current_table = []
    in_table = False
    font_family = config.get('document_settings', {}).get('font_family', 'Aptos')
    
    while i < len(lines):
        line = lines[i].strip()
        
        # Skip empty lines and markdown artifacts - reduce spacing
        if not line or line == '---' or re.match(r'^\s*[-*+]\s*$', line):
            # Only add paragraph break occasionally to reduce spacing
            if not in_table and i > 0 and i % 3 == 0:  # Reduced frequency
                p = doc.add_paragraph()
                p.paragraph_format.space_before = Pt(0)
                p.paragraph_format.space_after = Pt(1)
            i += 1
            continue
        
        # Handle tables
        if line.startswith('|') and '|' in line[1:]:
            if not in_table:
                in_table = True
                current_table = []
            
            # Skip separator lines
            if not re.match(r'^\|[\s\-\|]+\|$', line):
                current_table.append(line)
            
            i += 1
            continue
        
        # End of table
        if in_table:
            if current_table:
                table_data = parse_markdown_table('\n'.join(current_table))
                if table_data:
                    # Determine table type from previous heading
                    table_type = "performance" if "performance" in str(current_table).lower() else "cross_wp" if "cross" in str(current_table).lower() else "default"
                    create_formatted_table(doc, table_data, config, table_type)
                    # Minimal spacing after table
                    p = doc.add_paragraph()
                    p.paragraph_format.space_before = Pt(0)
                    p.paragraph_format.space_after = Pt(2)
            
            in_table = False
            current_table = []
        
        # Handle headings
        if line.startswith('#'):
            level = len(line) - len(line.lstrip('#'))
            heading_text = clean_text_content(line.lstrip('#').strip())
            
            if level == 1:
                if doc.paragraphs:  # Not the first heading
                    doc.add_page_break()
                p = doc.add_heading(heading_text, 0)
                p.style = 'Title Custom'
            elif level == 2:
                p = doc.add_heading(heading_text, 1)
                p.style = 'Heading 1 Custom'
            elif level == 3:
                p = doc.add_heading(heading_text, 2)
                p.style = 'Heading 2 Custom'
            elif level == 4:
                p = doc.add_heading(heading_text, 3)
                p.style = 'Heading 3 Custom'
            else:
                p = doc.add_paragraph(heading_text)
                p.style = 'Heading 3 Custom'
        
        # Handle metadata lines (lines with **key**: value)
        elif re.match(r'\*\*[^*]+\*\*:', line):
            p = doc.add_paragraph()
            # Parse the metadata
            parts = re.split(r'(\*\*[^*]+\*\*:)', line)
            for part in parts:
                if part.startswith('**') and part.endswith('**:'):
                    # Bold key
                    key = part[2:-3] + ':'
                    run = p.add_run(key)
                    run.bold = True
                elif part.strip():
                    # Regular value
                    p.add_run(' ' + part.strip())
        
        # Handle horizontal rules
        elif line.startswith('---'):
            p = doc.add_paragraph()
            p.add_run('_' * 50).italic = True
            p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        
        # Handle bullet points
        elif line.startswith('- ') or line.startswith('* '):
            bullet_text = line[2:].strip()
            
            # Handle bold text in bullets
            if '**' in bullet_text:
                p = doc.add_paragraph(style='List Bullet')
                parts = re.split(r'(\*\*[^*]+\*\*)', bullet_text)
                for part in parts:
                    if part.startswith('**') and part.endswith('**'):
                        run = p.add_run(part[2:-2])
                        run.bold = True
                    elif part:
                        p.add_run(part)
            else:
                doc.add_paragraph(bullet_text, style='List Bullet')
        
        # Handle numbered lists
        elif re.match(r'^\d+\.', line):
            list_text = re.sub(r'^\d+\.\s*', '', line)
            doc.add_paragraph(list_text, style='List Number')
        
        # Handle regular paragraphs
        elif line:
            # Always check for formatting - any asterisk needs processing
            if '*' in line:
                p = doc.add_paragraph()
                p.paragraph_format.space_before = Pt(0)
                p.paragraph_format.space_after = Pt(1)
                p.paragraph_format.line_spacing = 1.0
                body_font_size = config.get('text_styles', {}).get('body_text', {}).get('font_size', config.get('document_settings', {}).get('font_size', 10))
                
                # First handle bold formatting (**text**)
                processed_line = line
                bold_parts = re.split(r'(\*\*[^*]*\*\*)', processed_line)
                
                # Then handle italic formatting (*text*)
                final_parts = []
                for bold_part in bold_parts:
                    if not (bold_part.startswith('**') and bold_part.endswith('**')):
                        # Split non-bold parts by italic markers
                        italic_parts = re.split(r'(\*[^*]*\*)', bold_part)
                        final_parts.extend(italic_parts)
                    else:
                        final_parts.append(bold_part)
                
                # Process all parts and add to paragraph
                for part in final_parts:
                    clean_part = clean_text_content(part)
                    if part.startswith('**') and part.endswith('**'):
                        # Bold text
                        bold_text = clean_text_content(part[2:-2])
                        if bold_text:
                            run = p.add_run(bold_text)
                            run.bold = True
                            run.font.size = Pt(body_font_size)
                            run.font.name = font_family
                    elif part.startswith('*') and part.endswith('*') and not part.startswith('**'):
                        # Italic text
                        italic_text = clean_text_content(part[1:-1])
                        if italic_text:
                            run = p.add_run(italic_text)
                            run.italic = True
                            run.font.size = Pt(body_font_size)
                            run.font.name = font_family
                    elif clean_part:
                        # Regular text
                        run = p.add_run(clean_part)
                        run.font.size = Pt(body_font_size)
                        run.font.name = font_family
            else:
                clean_line = clean_text_content(line)
                if clean_line:  # Only add if not empty after cleaning
                    p = doc.add_paragraph(clean_line)
                    p.paragraph_format.space_before = Pt(0)
                    p.paragraph_format.space_after = Pt(1)
                    p.paragraph_format.line_spacing = 1.0
                body_font_size = config.get('text_styles', {}).get('body_text', {}).get('font_size', config.get('document_settings', {}).get('font_size', 10))
                for run in p.runs:
                    run.font.size = Pt(body_font_size)
                    run.font.name = font_family
        
        i += 1
    
    # Handle any remaining table
    if in_table and current_table:
        table_data = parse_markdown_table('\n'.join(current_table))
        if table_data:
            table_type = "performance" if "performance" in str(current_table).lower() else "cross_wp" if "cross" in str(current_table).lower() else "default"
            create_formatted_table(doc, table_data, config, table_type)

def convert_markdown_to_word(markdown_file, output_dir=None, filename_base=None):
    """Convert complete markdown analysis file to Word document"""
    if not os.path.exists(markdown_file):
        print(f"❌ Error: File not found - {markdown_file}")
        return None
    
    # Load configuration
    config = load_formatting_config()
    
    # Read markdown content
    with open(markdown_file, 'r', encoding='utf-8') as f:
        content = f.read()
    
    # Create Word document
    doc = Document()
    
    # Setup custom styles
    setup_document_styles(doc, config)
    
    # Set document margins from YAML configuration
    margin_inches = config.get('document_settings', {}).get('margin_inches', 1.0)
    sections = doc.sections
    for section in sections:
        section.top_margin = Inches(margin_inches)
        section.bottom_margin = Inches(margin_inches)
        section.left_margin = Inches(margin_inches)
        section.right_margin = Inches(margin_inches)
    
    # Process the complete markdown content
    process_markdown_content(content, doc, config)
    
    # Generate output filename
    if output_dir and filename_base:
        output_file = os.path.join(output_dir, f"{filename_base}.docx")
    else:
        base_name = os.path.splitext(os.path.basename(markdown_file))[0]
        output_file = f"output/{base_name}.docx"
    
    # Ensure output directory exists
    os.makedirs(os.path.dirname(output_file), exist_ok=True)
    
    # Save document
    doc.save(output_file)
    print(f"✅ Complete Word document created: {output_file}")
    
    # Get file size for confirmation
    file_size = os.path.getsize(output_file)
    print(f"📄 Document size: {file_size:,} bytes")
    
    return output_file

def main():
    """Main function for command line usage"""
    if len(sys.argv) != 2:
        print("Usage: python3 export_to_word.py <markdown_file>")
        print("Example: python3 export_to_word.py output/step3_wp_summary_v1.0_20250903.md")
        return
    
    markdown_file = sys.argv[1]
    convert_markdown_to_word(markdown_file)

if __name__ == "__main__":
    main()